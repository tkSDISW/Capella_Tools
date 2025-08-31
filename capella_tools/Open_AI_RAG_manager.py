import sys
from pathlib import Path
from openai import OpenAI
from IPython.display import display, Markdown
from IPython.core.display import HTML
from ipywidgets import widgets
import json
import os
from datetime import datetime
from bs4 import BeautifulSoup
import asyncio
from ipywidgets import widgets
from IPython.display import display
from IPython.core.display import Javascript
from jupyter_ui_poll import ui_events
import time
import ast
import re
import networkx as nx
import plotly.graph_objects as go
from pyvis.network import Network
import nbformat
from docx import Document
import PyPDF2
from capella_tools.model_configurator import get_api_key, get_base_url, get_model
Network(notebook=True)
import traceback



class ChatGPTAnalyzer:
    ALLOWED_EXTENSIONS = {'.yaml', '.yml', '.txt', '.xml', '.json', '.html', '.pdf', '.docx', '.sysml'}
    TEXT_BASED_EXTS = {'.yaml', '.yml', '.txt', '.xml', '.json', '.html', '.sysml'}
    PDF_EXTS = {'.pdf'}
    DOCX_EXTS = {'.docx'}

    def __init__(self, yaml_content=None, model=None, base_url=None, api_key=None, config_name=None):


        config = {}
        if config_name:
            config_path = Path.home() / ".secrets" / "model_configs.json"
            if config_path.exists():
                with config_path.open() as f:
                    configs = json.load(f)
                config = configs.get(config_name, {})
                if not config:
                    raise ValueError(f"No config named '{config_name}' found in model_configs.json.")
        elif model is None and base_url is None and api_key is None:
            # Use default config if no overrides and no name given
            config_path = Path.home() / ".secrets" / "model_configs.json"
            if config_path.exists():
                with config_path.open() as f:
                    configs = json.load(f)
                default_name = configs.get("_default")
                if default_name:
                    config = configs.get(default_name, {})
        
        # Choose from: passed args > config file > .secrets fallback
        self.api_key = api_key or config.get("api_key") or get_api_key()
        self.llm_url = base_url or config.get("base_url") or get_base_url()
        self.llm_model = model or config.get("model") or get_model()
    
        print(f"✅ ChatGPTAnalyzer initialized")
        print(f"🔐 API Key: {'Provided' if api_key else 'Loaded from secrets'}")
        print(f"🌐 Base URL: {self.llm_url or 'Default'}")
        print(f"🤖 Model: {self.llm_model}")
        self.yaml_content = yaml_content or ""
        self.chat_active = True
        self.messages = []
        if self.yaml_content:
            self.messages.append({
                "role": "system",
                "content": f"You are an expert in analyzing YAML files for system design. Here is the YAML file:\n---\n{self.yaml_content}\n---",
            })


    def submit_prompt(self, user_prompt, is_initial=True):
        user_prompt = user_prompt + " Format the response in .html format." if is_initial else user_prompt
        self.messages.append({"role": "user", "content": user_prompt})
        display(Markdown(f"**Your prompt:** {user_prompt}"))

    initial_prompt = lambda self, prompt: self.submit_prompt(prompt, is_initial=True)
    follow_up_prompt = lambda self, prompt: self.submit_prompt(prompt, is_initial=False)

    def add_text_file_to_messages(self, filepath):
        ext = os.path.splitext(filepath)[1].lower()
        if ext not in self.ALLOWED_EXTENSIONS:
            raise ValueError(f"Unsupported file type '{ext}'. Allowed types are: {', '.join(self.ALLOWED_EXTENSIONS)}")

        if ext in self.TEXT_BASED_EXTS:
            with open(filepath, 'r', encoding='utf-8') as f:
                content = f.read()
        elif ext in self.PDF_EXTS:
            with open(filepath, 'rb') as f:
                reader = PyPDF2.PdfReader(f)
                content = "\n".join(page.extract_text() for page in reader.pages if page.extract_text())
        elif ext in self.DOCX_EXTS:
            doc = Document(filepath)
            content = "\n".join([p.text for p in doc.paragraphs])
        else:
            content = ""

        self.messages.append({
            "role": "user",
            "content": f"File `{filepath}` was added for analysis:\n---\n{content}\n---"
        })
        print(f"✅ File `{filepath}` added to messages for analysis.")
        
    def get_response(self):
        """Send messages to ChatGPT and get a response with separate token usage info."""
        try:
            client = OpenAI(api_key=self.api_key, base_url=self.llm_url )
            response = client.chat.completions.create(
                messages=self.messages,
                model=self.llm_model,
            )
            assistant_message = response.choices[0].message.content
            usage = response.usage
            token_usage_info = (
                f"Tokens used: prompt={usage.prompt_tokens}, "
                f"completion={usage.completion_tokens}, total={usage.total_tokens}"
            ) if usage else "Token usage unavailable."
    
            self.messages.append({"role": "assistant", "content": assistant_message})
            
            # Strip unwanted code fences
            if assistant_message.startswith("```html"):
                assistant_message = assistant_message[7:]
            if assistant_message.endswith("```"):
                assistant_message = assistant_message[:-3]
            if assistant_message.startswith("```python"):
                assistant_message = assistant_message[9:]
            # Clean up with BeautifulSoup

            soup = BeautifulSoup(assistant_message, "html.parser")
            for tag in soup(["script", "style"]):
                tag.decompose()
            assistant_message_cleaned = str(soup)
    
            # Display response and then token usage separately
            if "<table" in assistant_message_cleaned or "<html" in assistant_message_cleaned:
                display(HTML(assistant_message_cleaned))
            else:
                display(Markdown(f"**Response:**\n\n{assistant_message_cleaned}"))
    
            # Display token info separately
            display(Markdown(f"**Token Usage Info:**\n\n{token_usage_info}"))
    
            return assistant_message_cleaned
    
        except Exception as e:
            error_type = type(e).__name__
            tb = traceback.format_exc()
        
            display(Markdown("❌ **Error communicating with the OpenAI-compatible API.**"))
            display(Markdown(f"**Exception:** `{error_type}`  \n**Message:** `{str(e)}`"))
        
            if "401" in str(e) or "Unauthorized" in str(e):
                display(Markdown("🔐 **It looks like your API key may be missing or invalid.**"))
            elif "403" in str(e) or "Permission" in str(e):
                display(Markdown("🚫 **You may not have access to this model or endpoint.**"))
            elif "Name or service not known" in str(e) or "Failed to establish a new connection" in str(e):
                display(Markdown("🌐 **Check your base URL. It may be incorrect or unreachable.**"))
            elif "model" in str(e).lower():
                display(Markdown("🤖 **There may be an issue with the model name.** Double-check your configuration."))

            # Optionally show traceback in collapsed view
            print("📄 Full Traceback (for debugging):")
            print(tb)

            return f"Error: {error_type} – {str(e)}"

        
    

    def generate_pyvis_graph_from_relations(self, relations, output_file="graph.html"):
        net = Network(height="750px", width="100%", directed=True, notebook=True, cdn_resources="in_line")
        net.set_options("""
        var options = {
          "physics": {
            "barnesHut": {
              "springLength": 90
            }
          },
          "edges": {
            "arrows": {"to": {"enabled": true}},
            "font": {"size": 14, "align": "middle"},
            "color": {"color": "gray"}
          }
        }
        """)
        added_nodes = set()
        for src, tgt, lbl in relations:
            for node in (src, tgt):
                if node not in added_nodes:
                    net.add_node(node, label=node, shape="ellipse")
                    added_nodes.add(node)
            net.add_edge(src, tgt, label=lbl)
        net.show(output_file)

    def analyze_and_generate_graph(self):
        special_prompt = """
Analyze the YAML content and extract relationships suitable for a graph.

- Each relationship should be a tuple: (source, target, label).
- Use simple labels like 'abstract type of','property value group','property value','constrains','linked model element',"region','states','transitions','outgoing transition','incoming transtion','do function','entry function','exits function','triggers','source state','destination state','after function','source','target','involves','exchange items','exchanges','allocated to', 'involving','components','deployed to','port','state machine','entities', 'elements of'.
- Use ref_id to navigate primary_id but do not list them in tuples.
- Output ONLY a Python list of these tuples.
- No explanation, no extra text, just the list in Python syntax.
"""
        print("Sending prompt to extract graphable relations...")
        self.submit_prompt(special_prompt)
        relations_text = self.get_response()
        try:
            print("Relations Text:", relations_text)
            cleaned_text = re.sub(r"^```[a-zA-Z]*\n", "", relations_text.strip())
            cleaned_text = cleaned_text.replace("```", "").strip()
            relations = ast.literal_eval(cleaned_text)
            if isinstance(relations, list) and all(isinstance(r, tuple) for r in relations):
                self.generate_pyvis_graph_from_relations(relations)
            else:
                print("Relations:", relations_text)
                raise ValueError("Extracted data is not a valid list of tuples.")
        except Exception as e:
            print(f"Failed to parse structured relations ({e}). Falling back to text summary.")
            fallback_prompt = """
Please summarize the key conclusions from the YAML content in a short bullet list.
Format it nicely for .html display.
"""
            self.submit_prompt(fallback_prompt)
            fallback_response = self.get_response()
            display(HTML(f"Fallback Summary:**\n\n{fallback_response}"))

    def interactive_chat(self):
        print("Starting interactive chat...")
        chat_history = widgets.Output()
        user_input = widgets.Textarea(
            placeholder="Type your prompt...",
            rows=3,
            layout=widgets.Layout(width="100%", border="2px solid #4A90E2", border_radius="8px",
                                  padding="12px", background_color="#F7F9FC", 
                                  box_shadow="3px 3px 10px rgba(0, 0, 0, 0.1)")
        )
        send_button = widgets.Button(description="Execute", button_style="primary")
        exit_button = widgets.Button(description="Exit", button_style="danger")
        file_list = [f for f in os.listdir(os.getcwd()) if os.path.isfile(f) and os.path.splitext(f)[1].lower() in self.ALLOWED_EXTENSIONS]
        file_dropdown = widgets.Dropdown(
            options=[""] + file_list,
            description="Load file:",
            layout=widgets.Layout(width="auto")
        )

        def load_file(_):
            filename = file_dropdown.value
            if not filename:
                return
            try:
                self.add_text_file_to_messages(filename)
                with chat_history:
                    display(Markdown(f"✅ **File `{filename}` was added and appended for analysis.**"))
            except Exception as e:
                with chat_history:
                    display(Markdown(f"❌ Error reading `{filename}`: {str(e)}"))

        file_dropdown.observe(load_file, names="value")

        def send_message(_):
            prompt = user_input.value.strip()
            if not prompt:
                return
            with chat_history:
                self.submit_prompt(prompt, is_initial=False)
                display(Markdown("**Generating a response..** "))
                chatbot_response = self.get_response()
            user_input.value = ""

        def exit_chat(_):
            self.chat_active = False

        send_button.on_click(send_message)
        exit_button.on_click(exit_chat)
        display(chat_history, user_input, widgets.HBox([send_button, exit_button]), file_dropdown)
        print("Waiting for chat interactions...")
        with ui_events() as poll:
            while self.chat_active:
                poll(10)
                time.sleep(1)



    def save_to_word(self, filename: str, html_response: str):
        """
        Save an HTML-formatted ChatGPT response to a Word document.
        
        Parameters:
        - filename: str - the name of the Word file to save (e.g., 'output.docx')
        - html_response: str - HTML content returned by the assistant
        """
        from docx import Document
        from bs4 import BeautifulSoup
        import os
    
        soup = BeautifulSoup(html_response, 'html.parser')
        doc = Document()
    
        for element in soup.find_all(['h1', 'h2', 'h3', 'p', 'li', 'ul', 'ol', 'strong']):
            text = element.get_text(strip=True)
    
            if element.name.startswith('h'):
                doc.add_heading(text, level=int(element.name[1]))
            elif element.name == 'p':
                doc.add_paragraph(text)
            elif element.name == 'li':
                doc.add_paragraph(text, style='ListBullet')
            elif element.name == 'strong':
                doc.add_paragraph(text).bold = True
    
        output_path = os.path.join(os.getcwd(), filename)
        doc.save(output_path)
        print(f"✅ Word document saved: {output_path}")

    def save_to_csv(self, filename: str, csv_response: str):
        """
        Save a CSV-formatted string to a .csv file for Excel compatibility.
    
        Parameters:
        - filename: str - the output file name (e.g., 'output.csv')
        - csv_response: str - response text that is CSV-formatted
        """
        import csv
        import os
        from io import StringIO
    
        # Use StringIO to read CSV string as file-like
        csv_stream = StringIO(csv_response.strip())
    
        reader = csv.reader(csv_stream)
        output_path = os.path.join(os.getcwd(), filename)
    
        with open(output_path, "w", newline='', encoding='utf-8') as f:
            writer = csv.writer(f)
            for row in reader:
                writer.writerow(row)
    
        print(f"✅ CSV file saved: {output_path}")